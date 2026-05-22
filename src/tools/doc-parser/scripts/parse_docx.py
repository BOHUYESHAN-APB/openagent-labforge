#!/usr/bin/env python3
"""
Standardized DOCX parser — extracts text with paragraph structure preserved.

Output: structured JSON with paragraphs, full_text, and optional style info.

Fallback parser for LLMs that cannot natively read DOCX files.
"""
import json
import sys
import argparse
from pathlib import Path


def extract_docx(docx_path: str) -> dict:
    """Extract paragraphs and structure from a DOCX file."""
    from docx import Document

    doc = Document(docx_path)
    paragraphs = []

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        entry = {
            "index": i,
            "text": text,
        }

        # Include style info if available
        if para.style and para.style.name:
            entry["style"] = para.style.name

        # Detect heading level
        style_name = para.style.name if para.style else ""
        if style_name and style_name.startswith("Heading"):
            level = style_name.replace("Heading ", "")
            try:
                entry["heading_level"] = int(level)
            except ValueError:
                pass

        paragraphs.append(entry)

    # Try to extract tables too
    tables = []
    for t_idx, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        tables.append({
            "index": t_idx,
            "rows": rows,
        })

    full_text = "\n".join(p["text"] for p in paragraphs)

    return {
        "filename": Path(docx_path).name,
        "paragraphs": paragraphs,
        "tables": tables if tables else None,
        "full_text": full_text,
    }


def main():
    parser = argparse.ArgumentParser(
        description="Extract structured text from DOCX"
    )
    parser.add_argument("docx", help="Path to DOCX file")
    parser.add_argument("--json", action="store_true",
                       help="Output structured JSON")
    parser.add_argument("-o", "--output",
                       help="Output file path (default: stdout)")
    args = parser.parse_args()

    if not Path(args.docx).exists():
        print(f"Error: file not found: {args.docx}", file=sys.stderr)
        sys.exit(1)

    try:
        result = extract_docx(args.docx)
        output = json.dumps(result, ensure_ascii=False, indent=2)

        if args.output:
            Path(args.output).write_text(output, encoding="utf-8")
            print(f"Written to: {args.output}", file=sys.stderr)
        else:
            print(output)

    except ImportError as e:
        print(f"Error: missing dependency. Install: pip install python-docx\n{e}",
              file=sys.stderr)
        sys.exit(1)
    except Exception as e:
        print(f"Error parsing DOCX: {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
